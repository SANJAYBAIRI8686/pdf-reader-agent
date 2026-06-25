import os
# pyright: ignore [reportMissingImports]
from pypdf import PdfReader
# pyright: ignore [reportMissingImports]
import docx
from app.core.logging import logger

class PDFParser:
    """
    Parser for PDF documents using pypdf.
    """
    @staticmethod
    def extract_text(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at {file_path}")
            
        reader = PdfReader(file_path)
        text_pages = []
        
        for idx, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_pages.append(page_text)
            except Exception as e:
                logger.warning("Failed to extract page text", file=file_path, page_idx=idx, error=str(e))
                continue
                
        return "\n\n".join(text_pages)

class DocxParser:
    """
    Parser for Microsoft Word (.docx) documents using python-docx.
    """
    @staticmethod
    def extract_text(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at {file_path}")
            
        doc = docx.Document(file_path)
        text_paragraphs = []
        
        for idx, para in enumerate(doc.paragraphs):
            # Capture paragraph text if not empty
            if para.text.strip():
                text_paragraphs.append(para.text)
                
        # Also capture text inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_paragraphs.append(cell.text)
                        
        return "\n\n".join(text_paragraphs)

class MarkdownParser:
    """
    Parser for Markdown (.md) and raw text files.
    """
    @staticmethod
    def extract_text(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at {file_path}")
            
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

class DocumentParser:
    """
    Central dispatcher factory for document parsing.
    """
    @classmethod
    def extract_text(cls, file_path: str, file_type: str) -> str:
        file_type = file_type.lower().strip(".")
        logger.info("Extracting document text...", file=file_path, type=file_type)
        
        if file_type == "pdf":
            return PDFParser.extract_text(file_path)
        elif file_type == "docx":
            return DocxParser.extract_text(file_path)
        elif file_type in ["md", "markdown", "txt"]:
            return MarkdownParser.extract_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_type}")
