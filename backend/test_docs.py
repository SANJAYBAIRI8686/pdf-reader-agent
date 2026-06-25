import os
import sys
import time
import docx
from pathlib import Path
from fastapi.testclient import TestClient

# Add current path to python search path
sys.path.append(str(Path(__file__).resolve().parent))

from app.main import app
from app.database.session import SessionLocal
from app.models.user import User
from app.models.document import Document

client = TestClient(app)

def test_document_workflow():
    print("\n=== STARTING DOCUMENT & BG-PARSING SUITE TESTS ===")
    
    test_email = "doc_junior_dev_tester@example.com"
    test_password = "SecurePassword123!"
    test_name = "Junior File Tester"
    
    # 1. Pre-test cleanup: ensure database is clean of this user
    db = SessionLocal()
    try:
        existing_user = db.query(User).filter(User.email == test_email).first()
        if existing_user:
            # Delete associated document vectors in ChromaDB first
            from app.rag.vectorstore import VectorStoreManager
            for doc in existing_user.documents:
                VectorStoreManager.delete_document_vectors(doc_id=doc.id)
            db.delete(existing_user)
            db.commit()
    finally:
        db.close()

    # 2. Register & Login to get token
    client.post(
        "/api/v1/auth/register",
        json={"email": test_email, "password": test_password, "full_name": test_name}
    )
    login_response = client.post(
        "/api/v1/auth/login",
        data={"username": test_email, "password": test_password}
    )
    token = login_response.json()["access_token"]
    auth_headers = {"Authorization": f"Bearer {token}"}

    # 3. Test extension validation (Upload an invalid format)
    print("Testing extension rejection (POST /upload with invalid format)...")
    invalid_file = {"file": ("script.exe", b"malicious-bytes-here", "application/octet-stream")}
    bad_upload = client.post("/api/v1/documents/upload", files=invalid_file, headers=auth_headers)
    assert bad_upload.status_code == 400
    assert "Unsupported file format" in bad_upload.json()["detail"]
    print("✓ Rejection of invalid extension passed.")

    # 4. Upload a valid Markdown file
    print("Testing Markdown upload...")
    md_content = b"## Topic: Artificial Intelligence\nMachine learning is a subset of AI. Agents perform tasks autonomously."
    md_file = {"file": ("ai_notes.md", md_content, "text/markdown")}
    upload1 = client.post("/api/v1/documents/upload", files=md_file, headers=auth_headers)
    assert upload1.status_code == 201
    upload1_json = upload1.json()
    assert upload1_json["is_duplicate"] is False
    doc1_id = upload1_json["document"]["id"]
    print("✓ Markdown upload succeeded.")

    # 5. Test duplicate detection (uploading same Markdown file again)
    print("Testing duplicate content detection...")
    md_file_dup = {"file": ("ai_notes_copy.md", md_content, "text/markdown")}
    upload2 = client.post("/api/v1/documents/upload", files=md_file_dup, headers=auth_headers)
    assert upload2.status_code == 200  # Returns status code 200 OK for duplicates!
    upload2_json = upload2.json()
    assert upload2_json["is_duplicate"] is True
    assert upload2_json["document"]["id"] == doc1_id  # Should point to the SAME document record
    print("✓ Duplicate upload protection passed.")

    # 6. Upload a programmatically generated Word Document (.docx)
    print("Programmatically creating a Word (.docx) file for testing parser...")
    docx_path = Path("temp_test_doc.docx")
    doc = docx.Document()
    doc.add_heading("Deep Learning Overview", level=1)
    doc.add_paragraph("Neural networks process layers of features to extract structural features.")
    doc.add_paragraph("Keywords: deeplearning, intelligence, networks, optimization, layers.")
    doc.save(docx_path)
    
    try:
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
            
        print("Testing Word (.docx) file upload and parser...")
        docx_file = {"file": ("dl_summary.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        upload3 = client.post("/api/v1/documents/upload", files=docx_file, headers=auth_headers)
        assert upload3.status_code == 201
        upload3_json = upload3.json()
        doc3_id = upload3_json["document"]["id"]
        print("✓ Word document upload succeeded.")
        
        # 7. Wait brief moment for background parsing thread to run and verify status updates
        print("Waiting 2 seconds for background parsing threads to complete...")
        time.sleep(2)
        
        # Fetch document status
        print("Verifying document status update (GET /documents/{doc_id})...")
        status1 = client.get(f"/api/v1/documents/{doc1_id}", headers=auth_headers)
        status3 = client.get(f"/api/v1/documents/{doc3_id}", headers=auth_headers)
        
        doc1_data = status1.json()
        doc3_data = status3.json()
        
        # Assert both finished successfully
        assert doc1_data["status"] == "processed", f"Doc 1 status: {doc1_data['status']}"
        assert doc3_data["status"] == "processed", f"Doc 3 status: {doc3_data['status']}"
        
        # Check generated summary and keywords
        assert "Artificial Intelligence" in doc1_data["summary"]
        assert doc3_data["keywords"] is not None
        print(f"✓ Background parsing validated. Doc 3 Top Keywords: {doc3_data['keywords']}")
        
        # 8. List documents upload check
        print("Testing document listing (GET /api/v1/documents/)...")
        list_response = client.get("/api/v1/documents/", headers=auth_headers)
        assert list_response.status_code == 200
        list_json = list_response.json()
        assert len(list_json) == 2  # Only doc1 and doc3 should be created
        print("✓ Document list query passed.")
        
        # 9. Test document delete
        print("Testing document deletion (DELETE /api/v1/documents/{doc_id})...")
        delete_response = client.delete(f"/api/v1/documents/{doc1_id}", headers=auth_headers)
        assert delete_response.status_code == 200
        
        # Ensure it is gone
        gone_response = client.get(f"/api/v1/documents/{doc1_id}", headers=auth_headers)
        assert gone_response.status_code == 404
        print("✓ Document deletion and storage cleanup passed.")
        
    finally:
        # Clean up local temporary test file
        if docx_path.exists():
            os.remove(docx_path)

    # 10. Post-test Database Cleanup
    print("Running final test database cleanup...")
    db = SessionLocal()
    try:
        user_record = db.query(User).filter(User.email == test_email).first()
        if user_record:
            # Delete associated document vectors in ChromaDB first
            from app.rag.vectorstore import VectorStoreManager
            for doc in user_record.documents:
                VectorStoreManager.delete_document_vectors(doc_id=doc.id)
            # Cascading deletes should remove remaining documents (doc3)
            db.delete(user_record)
            db.commit()
            print("✓ Database cleaned successfully.")
    except Exception as e:
        print(f"Database cleanup failure: {e}")
        db.rollback()
    finally:
        db.close()
        
    print("=== ALL DOCUMENT & BG-PARSING SUITE TESTS PASSED SUCCESSFULLY ===\n")

if __name__ == "__main__":
    test_document_workflow()
